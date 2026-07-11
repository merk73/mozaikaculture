from pathlib import Path
import re

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(r"C:\Users\Merk\Documents\Mozaika")
SOURCE = ROOT / "Мифология_коренных_народов_Хабаровского_края.docx"
OUT = ROOT / "Мифология_коренных_народов_Хабаровского_края_академический_текст.docx"


def set_tnr(run, size=14, bold=None, italic=None):
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    set_tnr(run, 12)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


src = Document(SOURCE)
chapter_titles = [
    "1. Нанайцы", "2. Ульчи", "3. Нивхи", "4. Орочи",
    "5. Удэгейцы", "6. Негидальцы", "7. Эвенки", "8. Эвены",
]
chapter_text = {title: [] for title in chapter_titles}
current = None
refs = []
in_bibliography = False

for p in src.paragraphs:
    text = p.text.strip()
    if text in chapter_titles:
        current = text
        in_bibliography = False
        continue
    if text == "Сравнительный обзор":
        current = None
    if text == "Библиография и электронные научные ресурсы":
        current = None
        in_bibliography = True
        continue
    if text == "Краткий словарь терминов":
        in_bibliography = False
        continue
    if current and text:
        # Keep only continuous academic prose; internal subheadings are omitted.
        if not p.style.name.startswith("Heading"):
            chapter_text[current].append(text)
    elif in_bibliography and text:
        if re.match(r"^\d+\.\s", text):
            refs.append(text)

doc = Document()
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(3.0)
section.right_margin = Cm(1.5)
section.header_distance = Cm(1.0)
section.footer_distance = Cm(1.0)

# Remove visible header text; retain only a centered page number in the footer.
section.header.paragraphs[0].text = ""
add_page_number(section.footer.paragraphs[0])

normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
normal.font.size = Pt(14)
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
normal.paragraph_format.first_line_indent = Cm(1.25)
normal.paragraph_format.line_spacing = 1.5
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(0)
normal.paragraph_format.widow_control = True

heading = doc.styles["Heading 1"]
heading.font.name = "Times New Roman"
heading._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
heading._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
heading._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
heading.font.size = Pt(14)
heading.font.bold = True
heading.font.color.rgb = None
heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
heading.paragraph_format.first_line_indent = Cm(0)
heading.paragraph_format.line_spacing = 1.5
heading.paragraph_format.space_before = Pt(12)
heading.paragraph_format.space_after = Pt(0)
heading.paragraph_format.keep_with_next = True

# Plain academic title, not a separate decorative cover.
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = Cm(0)
p.paragraph_format.line_spacing = 1.5
p.paragraph_format.space_after = Pt(12)
r = p.add_run("МИФОЛОГИЯ КОРЕННЫХ МАЛОЧИСЛЕННЫХ НАРОДОВ ХАБАРОВСКОГО КРАЯ")
set_tnr(r, 14, bold=True)

for title in chapter_titles:
    h = doc.add_paragraph(style="Heading 1")
    h.add_run(title)
    for text in chapter_text[title]:
        p = doc.add_paragraph(text)
        p.paragraph_format.keep_together = False

h = doc.add_paragraph(style="Heading 1")
h.add_run("Библиография")
for ref in refs:
    p = doc.add_paragraph(ref)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(-0.75)
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)

doc.core_properties.title = "Мифология коренных малочисленных народов Хабаровского края"
doc.core_properties.subject = "Академический реферат"
doc.core_properties.author = ""
doc.core_properties.keywords = "мифология, Хабаровский край, коренные малочисленные народы"
doc.save(OUT)
print(OUT)
print(f"chapters={len(chapter_text)}, paragraphs={sum(map(len, chapter_text.values()))}, references={len(refs)}")
