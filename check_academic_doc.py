from docx import Document
import sys

p = sys.argv[1] if len(sys.argv) > 1 else r"C:\Users\Merk\Documents\Mozaika\Мифология_коренных_народов_Хабаровского_края_академический_текст.docx"
d = Document(p)
heads = [x.text for x in d.paragraphs if x.style.name == "Heading 1"]
body = [x for x in d.paragraphs if x.style.name == "Normal" and x.text.strip()]
print("headings=", heads)
print("body_paragraphs=", len(body))
print("normal_font=", d.styles["Normal"].font.name, d.styles["Normal"].font.size.pt)
print("line_spacing=", d.styles["Normal"].paragraph_format.line_spacing)
print("first_line_cm=", round(d.styles["Normal"].paragraph_format.first_line_indent.cm, 2))
print("page_cm=", round(d.sections[0].page_width.cm, 1), round(d.sections[0].page_height.cm, 1))
print("margins_cm=", *(round(x.cm, 1) for x in (d.sections[0].left_margin, d.sections[0].right_margin, d.sections[0].top_margin, d.sections[0].bottom_margin)))
